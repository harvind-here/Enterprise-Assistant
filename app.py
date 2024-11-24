from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
import os
import tempfile
from dotenv import load_dotenv
#from rag import Rag
from chatbot import Chatbot
from bad_language_filter import filter_bad_language
from security import validate_passcode, generate_hashed_passcode, is_passcode_set, sanitize_input
import io
from PyPDF2 import PdfReader
from docx import Document
from werkzeug.utils import secure_filename
import time
import shutil

load_dotenv()

app = Flask(__name__, static_folder='static')
CORS(app)

# Use a temporary directory for file uploads
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()

try:
    chatbot = Chatbot(api_key=os.getenv('GROQ_API_KEY'))
except Exception as e:
    print(f"Error initializing Chatbot: {str(e)}")
    # You might want to exit the application or handle this error appropriately

if not os.getenv('GROQ_API_KEY'):
    raise ValueError("GROQ_API_KEY is not set in the environment variables")

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/validate_passcode', methods=['POST'])
def validate_pass():
    data = request.json
    passcode = data.get('passcode')
    
    if not is_passcode_set():
        # Generate new passcode
        generate_hashed_passcode(passcode)
        return jsonify({"valid": True, "message": "New passcode set successfully"}), 200
    
    if validate_passcode(passcode):
        return jsonify({"valid": True}), 200
    else:
        return jsonify({"valid": False}), 401

@app.route('/chat', methods=['POST'])
def chat():
    passcode = request.form.get('passcode')
    if not validate_passcode(passcode):
        return jsonify({"error": "Invalid passcode"}), 401

    print("Passcode validated, entering chat function")
    user_input = sanitize_input(request.form.get('message', ''))
    files = request.files.getlist('files')
    
    print(f"Received message: {user_input}")
    print(f"Number of files: {len(files)}")

    document_processing = False
    document_context = ""

    if files:
        document_processing = True
        total_words = 0

        for file in files:
            filename = secure_filename(file.filename)
            content = ''

            if filename.endswith('.txt'):
                content = file.read().decode('utf-8')
                file.seek(0)  # Reset file pointer
            elif filename.endswith('.docx'):
                doc = Document(io.BytesIO(file.read()))
                file.seek(0)  # Reset file pointer
                content = ' '.join([paragraph.text for paragraph in doc.paragraphs])
            elif filename.endswith('.pdf'):
                pdf = PdfReader(io.BytesIO(file.read()))
                file.seek(0)  # Reset file pointer
                for page in pdf.pages:
                    content += page.extract_text()
            else:
                return jsonify({'error': 'Unsupported file type'})

            total_words += len(content.split())
            document_context += content + "\n\n"

        if total_words > 1500:
            return jsonify({'error': f'Attached documents exceed 1500 words (current: {total_words})'})

    if filter_bad_language(user_input):
        response = 'Your message contains inappropriate language.'
    
    else:
        if document_processing:
            # Send initial status for document processing
            chatbot.set_document_context(sanitize_input(document_context))
            
            # Simulate processing time (remove in production)
            time.sleep(1)
        
        response = chatbot.get_response(user_input)

    return jsonify({'status': 'complete', 'response': response})

@app.route('/static/<path:path>')
def send_static(path):
    return send_from_directory('static', path)

@app.route('/check_word_count', methods=['POST'])
def check_word_count():
    total_words = 0
    files = request.files.getlist('files')
    
    for file in files:
        try:
            total_words += count_words_in_file(file)
        except Exception as e:
            return jsonify({'error': f'Error processing file {file.filename}: {str(e)}'}), 400
    
    return jsonify({'total_words': total_words})

@app.route('/check_passcode_status', methods=['GET'])
def check_passcode_status():
    return jsonify({"passcode_set": is_passcode_set()})

@app.route('/upload_dataset', methods=['POST'])
def upload_dataset():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        filename = secure_filename(file.filename)
        file_extension = os.path.splitext(filename)[1].lower()
        allowed_extensions = {'.csv', '.xlsx', '.json', '.txt', '.pdf', '.docx'}
        
        if file_extension not in allowed_extensions:
            return jsonify({'error': 'Unsupported file type'}), 400
        
        file_path = os.path.join(r'C:\Users\harvi\sih\dataset', filename)
        try:
            file.save(file_path)
            return jsonify({'message': 'File uploaded successfully'}), 200
        except Exception as e:
            return jsonify({'error': f'Error saving file: {str(e)}'}), 500

def count_words_in_file(file):
    filename = secure_filename(file.filename)
    content = ''

    if filename.endswith('.txt'):
        content = file.read().decode('utf-8')
        file.seek(0)  # Reset file pointer
    elif filename.endswith('.docx'):
        doc = Document(io.BytesIO(file.read()))
        file.seek(0)  # Reset file pointer
        content = ' '.join([paragraph.text for paragraph in doc.paragraphs])
    elif filename.endswith('.pdf'):
        pdf = PdfReader(io.BytesIO(file.read()))
        file.seek(0)  # Reset file pointer
        for page in pdf.pages:
            content += page.extract_text()
    else:
        return 0  # Unsupported file type

    return len(content.split())

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=3000, debug=True)